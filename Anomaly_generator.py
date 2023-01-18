from cgitb import text
from fileinput import close
from msilib.schema import Directory, Error, Font
from sqlite3 import Row
from xml.dom.expatbuilder import theDOMImplementation
import pandas as pd
import numpy as np
import os


from tkinter import *
from tkinter import messagebox
import tkinter.filedialog
from tkinter.messagebox import showinfo
from tkinter import filedialog as fd
from tkinter import Menu
from PIL import Image, ImageTk
import tkinter.font as tkFont

from docx import Document
from docx.shared import Mm
#from collections import defaultdict

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt




MA_report = []
MTBF_report = []
MTTR_report = []
PM_report = []



MER_label = ''
MTBF_label = ''
MTR_label = ''




window = Tk()
window.title('Anomaly Report Generator v01.2.2')


window.geometry('810x500')

thiess_img = Image.open('img/ThiessIcon4.png')
resized = thiess_img.resize((130,40))
test = ImageTk.PhotoImage(resized)

label_img = tkinter.Label(image=test)
label_img.image = test



label_img.grid(column=0,row=0,pady=0,padx=0)


'''

report attributes

'''




Week_name_label = Label(window, text='Week of report',width=18)
Site_name_Label = Label(window, text="Site name",width=18)
Week_name_entry = Entry(window,width=60)
Site_name_entry = Entry(window,width=60)

'''

Report Adjustment

'''


adjustment_frame = LabelFrame(window,text='Report Adjustment')
adjustment_frame.grid(column=0,row=11,padx=3,pady=5)

directory_to_save_frame = LabelFrame(window,text='Directory to save files')
directory_to_save_frame.grid(column=1,row=11,padx=3,pady=5,columnspan=2)

'''Scrolled down limit'''

downtime_limit_label = Label(adjustment_frame,text='Downtime Limit (hours)',width=18,bg='white')

downtime_limit_string = Spinbox(adjustment_frame,from_=0.0,values=(0.0, 10.0, 20.0, 30.0, 40.0, 50.0, 60.0, 70.0, 80.0, 90.0, 100.0), to=100.0, width=10, textvariable=IntVar())

pad_y_downtime_limit = 4
pad_x_downtime_limit = 4

downtime_limit_label.grid(column=0,row=5,padx=5,pady=5,rowspan=2)
downtime_limit_string.grid(column=1,row=5,ipadx=pad_x_downtime_limit,ipady=pad_y_downtime_limit,rowspan=2)


'''
Check box
'''

'''Critical Unit Only'''
chk_state_critical_unit = BooleanVar()
chk_state_critical_unit.set(False)
chk_critical = Checkbutton(adjustment_frame, text = 'Critical Unit Only', var = chk_state_critical_unit)
chk_critical.grid(column=0,row=3)

chk_state_pareto = BooleanVar()
chk_state_pareto.set(False)
chk_pareto = Checkbutton(adjustment_frame, text = 'Include Pareto Report', var = chk_state_pareto)
chk_pareto.grid(column=0,row=4)


def fitLabel(event):
    label = event.widget
    if not hasattr(label, "original_text"):
        
        '''preserve the original text so we can restore it if the widget grows.'''
        
        label.original_text = label.cget("text")

    font = tkFont.nametofont(label.cget("font"))
    text = label.original_text
    max_width = event.width
    actual_width = font.measure(text)
    if actual_width <= max_width:

        '''the original text fits; no need to add ellipsis'''
        
        label.configure(text=text)
    else:
        
        '''the original text won't fit. Keep shrinking until it does'''
        
        while actual_width > max_width and len(text) > 1:
            text = text[2:]
            actual_width = font.measure("..."+text)
        label.configure(text="..."+text)



def browse_for_file_MER():
    filepath = fd.askopenfilename(filetypes=(("xlsx", "*.xlsx"), ("all files", "*.*")))
    if len(filepath) == 0 :
        return
    else:
        global MERSummary, MER_label
        MER_label = Label(window, text=filepath, width=51,justify=RIGHT,bg='white')
        MER_label.grid(column=1,row=5)
        MERSummary = pd.read_excel(filepath)
        MER_label.bind("<Configure>", fitLabel)




def browse_for_file_EMPS():
    filepath = fd.askopenfilename(filetypes=(("xlsx", "*.xlsx"), ("all files", "*.*")))
    if len(filepath) == 0 :
        return
    else:
        global EMPS, EMPS_label
        EMPS_label = Label(window, text=filepath, width=50,justify=RIGHT,bg='white')
        EMPS_label.grid(column=1,row=6)
        EMPS = pd.read_excel(filepath)
        EMPS_label.bind("<Configure>", fitLabel)

def browse_for_file_PDTD():
    filepath = fd.askopenfilename(filetypes=(("xlsx", "*.xlsx"), ("all files", "*.*")))
    if len(filepath) == 0 :
        return
    else:
        global PDTD, PDTD_label
        PDTD_label = Label(window, text=filepath, width=50,justify=RIGHT,bg='white')
        PDTD_label.grid(column=1,row=7)
        PDTD = pd.read_excel(filepath)
        PDTD_label.bind("<Configure>",fitLabel)

def browse_for_file_fleetdesc():
    filepath = fd.askopenfilename(filetypes=(("xlsx", "*.xlsx"), ("all files", "*.*")))
    if len(filepath) == 0 :
        return
    else:
        global fleet_desc_root, fleetdesc_label
        fleetdesc_label = Label(window, text=filepath, width=50,justify=RIGHT,bg='white')
        fleetdesc_label.grid(column=1,row=8)
        fleet_desc_root = pd.ExcelFile(filepath)
        fleetdesc_label.bind("<Configure>",fitLabel)




filepath_dir_save = '...'
directory_label = Label(directory_to_save_frame,width=50,text=filepath_dir_save,bg='white')
directory_label.grid(column=2,row=0)
directory_label.bind("<Configure>",fitLabel)

def browse_for_save_location():
    global dir
    dir = fd.askdirectory()
    if len(dir)==0:
        return
    else:
        directory_label.config(text=dir)
        global save_directory
        save_directory = dir






def generate_docx():

    fleet_desc_raw = pd.read_excel(fleet_desc_root,'MMA')
    fleet_desc_pivot_pre_filter = pd.read_excel(fleet_desc_root,'pivotMMA')
    fleet_desc_exception = pd.read_excel(fleet_desc_root,'exception')

    downtime_limit = float(downtime_limit_string.get())
    nPeriod = Week_name_entry.get()
    site = Site_name_entry.get()
    for_only_critical_unit_boolean = chk_state_critical_unit.get()
    add_pareto_report = chk_state_pareto.get()

    filename = f'Anomaly report {site} {nPeriod}.docx'
    file_pareto_name = f'Pareto report {site} {nPeriod}.docx'


   



    '''filter for only critical unit'''
    fleet_desc_pivot=fleet_desc_pivot_pre_filter[fleet_desc_pivot_pre_filter['fleet desc'].notna()]

    fleet_desc_pivot.dropna(inplace=for_only_critical_unit_boolean)
    fleet_desc_raw.rename(columns={'Unit Number':'Unit'},inplace=True)
    fleetdesc = pd.merge(fleet_desc_raw,fleet_desc_pivot,on='Description',how='inner')
    fleet_desc_exception.rename(columns={'Unit Number':'Unit'},inplace=True)


    '''
    MER Processing

    '''

    MERSummary.drop(MERSummary.columns[[0,1]],axis=1,inplace=True)
    MERSummary.dropna(inplace=True)
    MERSummary.reset_index(drop=True,inplace=True)
    MERSummary.columns = MERSummary.iloc[0]
    MERSummary.drop(0,inplace=True)
    indexNames = MERSummary[(MERSummary['Unit']=='OEM Totals')|(MERSummary['Unit']=='Class Total')|(MERSummary['Unit']=='Report Total')].index
    MERSummary.drop(indexNames, inplace = True)

    '''PDTD Processing'''
    PDTD.dropna(thresh=3,inplace=True)
    PDTD.reset_index(drop=True,inplace=True)
    PDTD.columns = PDTD.iloc[0]
    PDTD.drop(0,inplace=True)
    PDTD.rename(columns = {'Unit No.' : 'Unit'},inplace=True)

    PDTD_complete1 = pd.merge(PDTD, fleetdesc, on='Unit', how='left')
    PDTD_complete2 = pd.merge(PDTD, fleet_desc_exception, on='Unit', how='left')

    PDTD_complete2['fleet desc']=PDTD_complete2['fleet desc'].combine_first(PDTD_complete1['fleet desc']).fillna(0).astype(str)
    PDTD_complete2.dropna(subset=['fleet desc'],inplace=True)
    PDTD_complete = PDTD_complete2
    PDTD_complete['Period\nEvent\nMaint.\nDuration'] =  PDTD_complete['Period\nEvent\nMaint.\nDuration'].apply(lambda x: round(x,2))
    PDTD_complete['Activity'].replace(',','',regex=True,inplace=True)
    '''
    EMPS Processing

    '''
    EMPS.drop(EMPS.columns[[0]],axis=1,inplace=True)
    EMPS.dropna(inplace=True)
    EMPS.reset_index(drop=True,inplace=True)
    EMPS.columns = EMPS.iloc[0]
    EMPS.drop(0,inplace=True)
    EMPS.rename(columns = {'Unit\nNumber':'Unit'},inplace=True)
    indexNames2 = EMPS[(EMPS['Unit']=='OEM Totals')|(EMPS['Unit']=='Class Total')|(EMPS['Unit']=='Report Total')].index
    EMPS.drop(indexNames2,inplace=True)

    '''Combine MER & EMPS'''
    Combined = pd.merge(MERSummary,EMPS, on='Unit', how='outer')
    Combined1 =  pd.merge(Combined, fleet_desc_exception, on='Unit', how='left')
    Combined2 = pd.merge(Combined, fleetdesc, on='Unit', how='left')
    Combined1['fleet desc'] = Combined1['fleet desc'].combine_first(Combined2['fleet desc'])#.fillna(0).astype(str)
    Combined1.dropna(subset=['fleet desc'],inplace=True)
    All_KPI = Combined1

    '''Combine fleet desc & combined'''
    All_KPI.dropna(thresh=24, inplace=True)


    '''round selected column'''
    All_KPI['PMD\n(08)'] = All_KPI['PMD\n(08)'].apply(lambda x: round(x,2))
    All_KPI['UMD\n(09)'] = All_KPI['UMD\n(09)'].apply(lambda x: round(x,2))
    All_KPI['Total\nEngine On\nPeriod'] =  All_KPI['Total\nEngine On\nPeriod'].apply(lambda x: round(x,2))
    All_KPI['PDAM\n08020\n09020'] =  All_KPI['PDAM\n08020\n09020'].apply(lambda x: round(x,2))

    '''Gather all required variables'''
    PWT = All_KPI.groupby(['fleet desc'])['PWT\n(00)'].sum() #F
    SWT = All_KPI.groupby(['fleet desc'])['SWT\n(01)'].sum() #G
    IODOn = All_KPI.groupby(['fleet desc'])['IOD-On\n(04)'].sum() #H
    IODOff = All_KPI.groupby(['fleet desc'])['IOD-Off\n(04)'].sum() #I
    EODOn = All_KPI.groupby(['fleet desc'])['EOD-On\n(05)'].sum() #J
    EODOff = All_KPI.groupby(['fleet desc'])['EOD-Off\n(05)'].sum() #K
    PDAM = All_KPI.groupby(['fleet desc'])['PDAM\n08020\n09020'].sum() #T
    PMD = All_KPI.groupby(['fleet desc'])['PMD\n(08)'].sum() #L
    UMD = All_KPI.groupby(['fleet desc'])['UMD\n(09)'].sum() #M
    POMR = All_KPI.groupby(['fleet desc'])['POMR \n&\nPOMN\n(18+28)'].sum()#Q
    UOMR = All_KPI.groupby(['fleet desc'])['UOMR\n&\nUOMN\n(19+29)'].sum()#R
    EngOn = All_KPI.groupby(['fleet desc'])['Total\nEng On\n(SMU Hrs)'].sum() #V
    nFail = All_KPI.groupby(['fleet desc'])['Number of\nFailures\nPeriod'].sum() #W
    nFailsum = All_KPI.groupby(['fleet desc'])['Number of\nFailures\nPeriod'].sum()
    MTBF_period = All_KPI.groupby(['fleet desc'])['MTBF\nPeriod'].sum() #X
    MTTRF_period = All_KPI.groupby(['fleet desc'])[['MTTR-F\nPeriod','Number of\nFailures\nPeriod']] #Y

    All_KPI['Failure Hours'] = All_KPI['Number of\nFailures\nPeriod']*All_KPI['MTTR-F\nPeriod']
    fail_hours = All_KPI.groupby(['fleet desc'])['Failure Hours'].sum()


    '''Do the math
    
    -- new branch to add pm
    '''
    MA = ((PWT+SWT+IODOn+IODOff+EODOn+EODOff+PDAM)/(PWT+SWT+IODOn+IODOff+EODOn+EODOff+PDAM+PMD+UMD-PDAM).replace({0:np.nan})).fillna(0)
    MTBF =((EngOn)/(nFail).replace({0:np.nan})).fillna(EngOn)
    MTTR = ((fail_hours)/(nFailsum).replace({0:np.nan})).fillna(fail_hours)
    PM = ((PMD+POMR)/(PMD+UMD+POMR+UOMR-PDAM).replace({0:np.nan})).fillna(0)

    '''Combined all KPI'''
    Anomali_db = pd.merge(MA.to_frame(),MTBF.to_frame(), on='fleet desc', how= 'outer').merge(MTTR.to_frame(), on='fleet desc', how= 'outer')
    #Anomali_db = pd.merge(Anomali_db1,MTTR.to_frame(), on='fleet desc', how= 'outer')
    complete_anomali = pd.merge(Anomali_db, PM.to_frame(), on='fleet desc', how= 'outer')
    #complete_anomali.rename(columns = {'0_x':'MA','0_y':'MTBF'},inplace=True)
    complete_anomali.columns = ['MA','MTBF','MTTR','PM']
    '''round by 2 decimals'''
    complete_anomali['MTBF'] = complete_anomali['MTBF'].apply(lambda x: round(x,2))
    complete_anomali['MTTR'] = complete_anomali['MTTR'].apply(lambda x: round(x,2))
    '''drop zeros'''
    complete_KPI = complete_anomali.loc[~(complete_anomali==0).all(axis=1)]

    combined_MTBF_standard = pd.merge(complete_KPI,fleet_desc_pivot.groupby('fleet desc').mean('MTBF Standard'),how='left',on='fleet desc')

    red_MTBF_this_week1 = combined_MTBF_standard[combined_MTBF_standard['MTBF']<combined_MTBF_standard['MTBF Standard']]#.dropna().set_index('fleet desc')['MTBF']
    red_MTBF_this_week = red_MTBF_this_week1['MTBF']




    'Processing MTBF'

    MTBF_df = PDTD_complete[PDTD_complete['fleet desc'].isin(red_MTBF_this_week.index)]
    '''MTBF Report template'''
    MTBF_report = []


    for i in red_MTBF_this_week.index:
        comp_gp_issue = []
        
        '''
        filter downtime detail for only failure event 
        '''
        
        na = MTBF_df[(MTBF_df['Is Failure']=='Yes')&(MTBF_df['fleet desc'] ==i)][['Reason','System']].groupby(['Reason'],as_index=False).count().sort_values('System',ascending=False)['Reason']
        count_of_failure = MTBF_df[(MTBF_df['fleet desc'] ==i)&(MTBF_df['Is Failure']=='Yes')].count()['Is Failure']
        '''
        conditional statement for unit that has no failure
        '''
        if count_of_failure == 0 or count_of_failure == 1:
            continue
        else:
            for u in na:    
                '''
                Create dataframe for system failure
                '''
                system_failure = MTBF_df[(MTBF_df['Is Failure']=='Yes')][['fleet desc','Reason','System','Part']].groupby(['fleet desc','Reason','System'],as_index=False).count()
                system_failure['MTBF_gp_repetitive'] = system_failure['System']+' '+system_failure['Part'].values.astype(str)+' times'
                
                '''
                Create dataframe for reason failure failure
                '''
                reason_failure = MTBF_df[(MTBF_df['fleet desc']==i) & (MTBF_df['Is Failure']=='Yes')]['Reason'].value_counts().reset_index()
                reason_failure['Reason_repetitive'] = reason_failure['index']+' '+reason_failure['Reason'].values.astype(str)+' times'

                general_comp_failure = []
                for e in reason_failure['index']:
                    f = system_failure[(system_failure['fleet desc']==i)&(system_failure['Reason']==e)]['MTBF_gp_repetitive'].tolist()
                    general_comp_failure.append(f)
                repetitive_issue = dict(zip(reason_failure['Reason_repetitive'],general_comp_failure))
            comp_gp_issue.append(repetitive_issue)
        '''
        report template per fleet model
        
        '''
        
        structured_MTBF = {
        'Fleet':i,
        'MTBF':red_MTBF_this_week[i],
        'Total Unit':All_KPI.groupby('fleet desc').count()['Unit'][i],
        'Total Hours Engine ON' : All_KPI.groupby('fleet desc').sum()['Total\nEngine On\nPeriod'][i],
        'Count of failures' : count_of_failure,
        'Failures by Comp Gp with repetitive issue' : comp_gp_issue,
        'MTBF Standard' : combined_MTBF_standard.groupby('fleet desc').mean().dropna().loc[i]['MTBF Standard']
        }
        
        MTBF_report.append(structured_MTBF)

        
        
    'Processing MA'


    red_MA_this_week = complete_KPI['MA'][complete_KPI['MA']<0.9].dropna()
    MA_df= PDTD_complete.loc[PDTD_complete['fleet desc'].isin(red_MA_this_week.index)]


    desc_MA= MA_df.loc[:,('Unit')]+' '+MA_df.loc[:,('Reported Fault/Job Description')]+' ('+MA_df.loc[:,('Period\nEvent\nMaint.\nDuration')].values.astype(str)+" Hours)"
    MA_df.groupby(['Unit','Event\nID','Period\nEvent\nMaint.\nDuration'])['Reported Fault/Job Description'].apply(lambda x: ', '.join(x.astype(str))).reset_index()
    MA_df_new = MA_df.merge(desc_MA.rename('MA Anomali'),left_index=True, right_index=True).sort_values(['Period\nEvent\nMaint.\nDuration'],ascending=False)
    MA_report_fail_df = MA_df.groupby(['Unit','Event\nID','Activity','fleet desc','Period\nEvent\nMaint.\nDuration'])['Reported Fault/Job Description'].apply(lambda x: ','.join(x.astype(str))).reset_index()
    MA_report_fail_df['MA Anomali'] = MA_report_fail_df.loc[:,('Unit')]+' '+MA_report_fail_df.loc[:,('Reported Fault/Job Description')]+' ('+MA_report_fail_df.loc[:,('Period\nEvent\nMaint.\nDuration')].values.astype(str)+" Hours)"


    '''MA Report Template'''

    MA_report = []

        
        
    for i in red_MA_this_week.index:
        Scheduled_maintenance = []
        Unschedule_maintenance = []
        
        '''Processing MA df scheduled & unscheduled''' 
        
        MA_df_sched = MA_report_fail_df[((MA_report_fail_df.Activity =='08-Planned Maintenance (PMD)')&(MA_report_fail_df['fleet desc'] ==i)&(MA_report_fail_df['Period\nEvent\nMaint.\nDuration']>downtime_limit))].sort_values(['Period\nEvent\nMaint.\nDuration'],ascending=False)
        MA_df_unsched = MA_report_fail_df[((MA_report_fail_df.Activity =='09-Unplanned Maintenance (UMD)')&(MA_report_fail_df['fleet desc'] ==i)&(MA_report_fail_df['Period\nEvent\nMaint.\nDuration']>downtime_limit))].sort_values(['Period\nEvent\nMaint.\nDuration'],ascending=False)
        
        Scheduled_fleet_df = MA_df_sched['MA Anomali']
        Unscheduled_fleet_df = MA_df_unsched['MA Anomali'] 
        
        for x in Scheduled_fleet_df:
            Scheduled_maintenance.append(x)
        
        for x in Unscheduled_fleet_df:
            Unschedule_maintenance.append(x)
        
        
        structured_MA = {
            
        
        'Fleet':i,
        'MA':red_MA_this_week[i],
        'Total Unit':All_KPI.groupby('fleet desc').count()['Unit'][i],
        'UMD total hours':round(All_KPI.groupby('fleet desc').sum()['UMD\n(09)'][i],2),
        'PMD total hours':All_KPI.groupby('fleet desc').sum()['PMD\n(08)'][i],
        'PDAM total hours':All_KPI.groupby('fleet desc').sum()['PDAM\n08020\n09020'][i],
        'Scheduled maintenance' : Scheduled_maintenance,
        'Unscheduled maintenance' : Unschedule_maintenance
        
    }
        MA_report.append(structured_MA)
    
    '''Processing PM'''


    red_PM_this_week = complete_KPI['PM'][complete_KPI['PM']<0.6].dropna()
    PM_df= PDTD_complete.loc[PDTD_complete['fleet desc'].isin(red_PM_this_week.index)]


    desc_PM= PM_df.loc[:,('Unit')]+' '+PM_df.loc[:,('Reported Fault/Job Description')]+' ('+PM_df.loc[:,('Period\nEvent\nMaint.\nDuration')].values.astype(str)+" Hours)"
    PM_df.groupby(['Unit','Event\nID','Period\nEvent\nMaint.\nDuration'])['Reported Fault/Job Description'].apply(lambda x: ', '.join(x.astype(str))).reset_index()
    PM_df_new = PM_df.merge(desc_MA.rename('PM Anomali'),left_index=True, right_index=True).sort_values(['Period\nEvent\nMaint.\nDuration'],ascending=False)
    PM_report_fail_df = PM_df.groupby(['Unit','Event\nID','Activity','fleet desc','Period\nEvent\nMaint.\nDuration'])['Reported Fault/Job Description'].apply(lambda x: ','.join(x.astype(str))).reset_index()
    PM_report_fail_df['PM Anomali'] = PM_report_fail_df.loc[:,('Unit')]+' '+PM_report_fail_df.loc[:,('Reported Fault/Job Description')]+' ('+PM_report_fail_df.loc[:,('Period\nEvent\nMaint.\nDuration')].values.astype(str)+" Hours)"


    '''PM Report Template'''

    PM_report = []

        
        
    for i in red_PM_this_week.index:
        Scheduled_maintenance = []
        Unschedule_maintenance = []
        
        '''Processing MA df scheduled & unscheduled''' 
        
        PM_df_sched = PM_report_fail_df[((PM_report_fail_df.Activity =='08-Planned Maintenance (PMD)')&(PM_report_fail_df['fleet desc'] ==i)&(PM_report_fail_df['Period\nEvent\nMaint.\nDuration']>downtime_limit))].sort_values(['Period\nEvent\nMaint.\nDuration'],ascending=False)
        PM_df_unsched = PM_report_fail_df[((PM_report_fail_df.Activity =='09-Unplanned Maintenance (UMD)')&(PM_report_fail_df['fleet desc'] ==i)&(PM_report_fail_df['Period\nEvent\nMaint.\nDuration']>downtime_limit))].sort_values(['Period\nEvent\nMaint.\nDuration'],ascending=False)
        
        Scheduled_fleet_df = PM_df_sched['PM Anomali']
        Unscheduled_fleet_df = PM_df_unsched['PM Anomali'] 
        
        for x in Scheduled_fleet_df:
            Scheduled_maintenance.append(x)
        
        for x in Unscheduled_fleet_df:
            Unschedule_maintenance.append(x)
        
        
        structured_PM = {
            
        
        'Fleet':i,
        'PM':red_PM_this_week[i],
        'Total Unit':All_KPI.groupby('fleet desc').count()['Unit'][i],
        'UMD total hours':round(All_KPI.groupby('fleet desc').sum()['UMD\n(09)'][i],2),
        'PMD total hours':All_KPI.groupby('fleet desc').sum()['PMD\n(08)'][i],
        'PDAM total hours':All_KPI.groupby('fleet desc').sum()['PDAM\n08020\n09020'][i],
        'Scheduled maintenance' : Scheduled_maintenance,
        'Unscheduled maintenance' : Unschedule_maintenance
        
    }
        PM_report.append(structured_PM)
        

    '''Processing MTTRF'''

    red_MTTR_this_week = complete_KPI['MTTR'][complete_KPI['MTTR']>6].dropna()
    MTTR_df = PDTD_complete[PDTD_complete['fleet desc'].isin(red_MTTR_this_week.index)]

    MTTRF_report_fail_df = MTTR_df.groupby(['Unit','Event\nID','Activity','fleet desc','Period\nEvent\nMaint.\nDuration','Is Failure'])['Description of Repair'].apply(lambda x: ','.join(x.astype(str))).reset_index()
    MTTRF_report_fail_df['MTTRF Anomali'] = MTTRF_report_fail_df.loc[:,('Unit')]+' '+MTTRF_report_fail_df.loc[:,('Description of Repair')]+' ('+MTTRF_report_fail_df.loc[:,('Period\nEvent\nMaint.\nDuration')].values.astype(str)+" Hours)"
    MTTRF_report_fail_df.replace({'MTTRF Anomali':{'\n':' '}},regex=True)

    '''MTTR Report Template'''

    MTTR_report = []
    
        
    for i in red_MTTR_this_week.index:
        Breakdown_detail = []
        
        MTTR_breakdown_fail_yes =  MTTRF_report_fail_df[(MTTRF_report_fail_df['Is Failure'] =='Yes')&(MTTRF_report_fail_df['fleet desc'] ==i)&(MTTRF_report_fail_df['Period\nEvent\nMaint.\nDuration']>downtime_limit)].sort_values(['Period\nEvent\nMaint.\nDuration'],ascending=False)
        
        for x in MTTR_breakdown_fail_yes['MTTRF Anomali']:
            Breakdown_detail.append(x)
            
        structured_MTTR = {
        'Fleet':i,
        'MTTR':red_MTTR_this_week[i],
        'Total Unit':All_KPI.groupby('fleet desc').count()['Unit'][i],
        'Hours Maintenance with Failure is Yes' : round(All_KPI[(All_KPI['fleet desc']==i)]['Failure Hours'].sum(),2),
        'Count of Failures' : MTTR_df[(MTTR_df['Is Failure'] =='Yes')&(MTTR_df['fleet desc'] ==i)].count()['Is Failure'],
        'Detail Breakdown with Failure is Yes' : Breakdown_detail
        
    }
        MTTR_report.append(structured_MTTR)
        

    '''Table size arrangement'''
        
        
    def set_col_widths(table):
        widths = (Mm(31.9), Mm(11.25),Mm(11.25), Mm(210.5))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                
    '''repeatable header'''
    def set_repeat_table_header(row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    '''
    Document Creator

    '''

    '''Create word document using word'''
    Anomali_report_full = Document()

    section = Anomali_report_full.sections[0]
    section.page_height = Mm(420)
    section.page_width = Mm(297)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(10.5)
    section.top_margin = Mm(22.2)
    section.bottom_margin = Mm(10.5)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)


    '''shading elm with parse_xml'''

    '''MA'''
    shading_elm1 = parse_xml(r'<w:shd {} w:fill="66FFFF"/>'.format(nsdecls('w')))
    shading_elm2 = parse_xml(r'<w:shd {} w:fill="66FFFF"/>'.format(nsdecls('w')))
    shading_elm3 = parse_xml(r'<w:shd {} w:fill="66FFFF"/>'.format(nsdecls('w')))
    shading_elm4 = parse_xml(r'<w:shd {} w:fill="66FFFF"/>'.format(nsdecls('w')))

    '''MTBF'''
    shading_elm5 = parse_xml(r'<w:shd {} w:fill="92D050"/>'.format(nsdecls('w')))
    shading_elm6 = parse_xml(r'<w:shd {} w:fill="92D050"/>'.format(nsdecls('w')))
    shading_elm7 = parse_xml(r'<w:shd {} w:fill="92D050"/>'.format(nsdecls('w')))
    shading_elm8 = parse_xml(r'<w:shd {} w:fill="92D050"/>'.format(nsdecls('w')))


    '''MTTR-F'''
    shading_elm9 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
    shading_elm10 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
    shading_elm11 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
    shading_elm12 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))

    '''PM'''
    shading_elm13 = parse_xml(r'<w:shd {} w:fill="1616F6"/>'.format(nsdecls('w')))
    shading_elm14 = parse_xml(r'<w:shd {} w:fill="1616F6"/>'.format(nsdecls('w')))
    shading_elm15 = parse_xml(r'<w:shd {} w:fill="1616F6"/>'.format(nsdecls('w')))
    shading_elm16 = parse_xml(r'<w:shd {} w:fill="1616F6"/>'.format(nsdecls('w')))




    '''Mechanical Availability'''
    Anomali_report_full.add_paragraph(text=f'Fleet with MA Below Target {site}').alignment = WD_ALIGN_PARAGRAPH.CENTER


    '''
    formatting MA table
    '''

    table_MA = Anomali_report_full.add_table(rows=1,cols=4)
    table_MA.style = 'Table Grid'



    hdr_cells_MA = table_MA.rows[0].cells
    for cell in hdr_cells_MA:
        cell.height = Mm(30)
        
    hdr_cells_MA[0].text = 'Fleet/Model'
    hdr_cells_MA[1].text = f'MA week {nPeriod}'
    hdr_cells_MA[2].text = 'Target'
    hdr_cells_MA[3].text = 'Comments on Non-Compliance'

    '''cells shading'''
    hdr_cells_MA[0]._tc.get_or_add_tcPr().append(shading_elm1)
    hdr_cells_MA[1]._tc.get_or_add_tcPr().append(shading_elm2)
    hdr_cells_MA[2]._tc.get_or_add_tcPr().append(shading_elm3)
    hdr_cells_MA[3]._tc.get_or_add_tcPr().append(shading_elm4)

    '''repeat header'''
    set_repeat_table_header(table_MA.rows[0])


    for i in MA_report:
        
        TOU_MA = i.get('Total Unit')
        PMD_MA = i.get('PMD total hours')
        UMD_MA = i.get('UMD total hours')
        PDAM_MA = i.get('PDAM total hours')
        
        row_cells_MA = table_MA.add_row().cells
        row_cells_MA[0].text = i.get('Fleet')
        row_cells_MA[1].text = "{:.2%}".format(i.get('MA'))
        row_cells_MA[2].text = "90%"
        
        sch = i.get('Scheduled maintenance')
        Unsch = i.get('Unscheduled maintenance')
        
        Scheduled_maintenance = '\n'.join([c for c in sch[1:]])
        Unscheduled_maintenance = '\n'.join([c for c in Unsch[1:]])
        
        non_compliance = f'Number of Units = {TOU_MA} \nPMD Total = {PMD_MA} \nUMD Total= {UMD_MA} \nPDAM Total hours= {PDAM_MA}\n'
        #MA_anomalies = f'Scheduled maintenance :\n{Scheduled_maintenance}\n\nUnscheduled maintenance :\n{Unscheduled_maintenance}'

        row_cells_MA[3].add_paragraph(text=non_compliance)
        row_cells_MA[3].add_paragraph('Scheduled maintenance :\n')
        for j in sch:
            row_cells_MA[3].add_paragraph(j,style='List Bullet')
        row_cells_MA[3].add_paragraph('\nUnscheduled maintenance :\n')
        for k in Unsch:
            row_cells_MA[3].add_paragraph(k,style='List Bullet')    

    set_col_widths(table_MA)
    Anomali_report_full.add_page_break()
        
    '''
    Mean Time Between Failure
    '''
    Anomali_report_full.add_paragraph(text=f'Fleet with MTBF Below Target {site}').alignment = WD_ALIGN_PARAGRAPH.CENTER

    '''
    formatting MTBF table 
    '''

    table_MTBF = Anomali_report_full.add_table(rows=1, cols=4)
    table_MTBF.style = 'Table Grid'

    hdr_cells_MTBF = table_MTBF.rows[0].cells
    hdr_cells_MTBF[0].text = 'Fleet/Model '
    hdr_cells_MTBF[1].text = f'MTBF week {nPeriod}'
    hdr_cells_MTBF[2].text = 'Target'
    hdr_cells_MTBF[3].text = 'Comments on Non-Compliance'

    '''cells shading'''
    hdr_cells_MTBF[0]._tc.get_or_add_tcPr().append(shading_elm5)
    hdr_cells_MTBF[1]._tc.get_or_add_tcPr().append(shading_elm6)
    hdr_cells_MTBF[2]._tc.get_or_add_tcPr().append(shading_elm7)
    hdr_cells_MTBF[3]._tc.get_or_add_tcPr().append(shading_elm8)

    '''repeat header'''
    set_repeat_table_header(table_MTBF.rows[0])


    for d in MTBF_report:
        
        
        
        TOU = d.get('Total Unit')
        THEON = d.get('Total Hours Engine ON')
        COF = d.get('Count of failures')
        
        
        
        MTBF_value = str(d.get('MTBF'))
        row_cells = table_MTBF.add_row().cells
        row_cells[0].text = d.get('Fleet')
        row_cells[1].text = f'MTBF={MTBF_value}'
        
        '''adding target MTBF'''
        row_cells[2].text = str(d.get('MTBF Standard'))
        
        gp_components = []
        
        
        for system in d.get('Failures by Comp Gp with repetitive issue'):
                for key, value in system.items():
                    s = ", ".join(value)
                    ter = f'{key}: {s}'
                    gp_components.append(ter)
                
                
        f_list = '\n'.join([d for d in gp_components[0:]])
        fleet_data = f'Number of Units = {TOU} \nTotal Hours Engine On = {THEON} \n Count of failures = {COF} \n Failures by Comp Gp with repetitive issue : \n'
        
        row_cells[3].add_paragraph(fleet_data)
        for j in gp_components:
            row_cells[3].add_paragraph(j,style='List Bullet')


        
    set_col_widths(table_MTBF)
    Anomali_report_full.add_page_break()
        
    '''
    Mean Time to Repair Failure
    '''
    Anomali_report_full.add_paragraph(text=f'Fleet with MTTR-F Below Target {site}').alignment = WD_ALIGN_PARAGRAPH.CENTER

    '''
    formatting MTTRF table 
    '''
    table_MTTR = Anomali_report_full.add_table(rows=1, cols=4)
    table_MTTR.style = 'Table Grid'

    hdr_cells_MTTR = table_MTTR.rows[0].cells
    hdr_cells_MTTR[0].text = 'Fleet/Model '
    hdr_cells_MTTR[1].text = f'MTTR week {nPeriod}'
    hdr_cells_MTTR[2].text = 'Target'
    hdr_cells_MTTR[3].text = 'Comments on Non-Compliance'

    '''cells shading'''
    hdr_cells_MTTR[0]._tc.get_or_add_tcPr().append(shading_elm9)
    hdr_cells_MTTR[1]._tc.get_or_add_tcPr().append(shading_elm10)
    hdr_cells_MTTR[2]._tc.get_or_add_tcPr().append(shading_elm11)
    hdr_cells_MTTR[3]._tc.get_or_add_tcPr().append(shading_elm12)

    '''repeat header'''
    set_repeat_table_header(table_MTTR.rows[0])

    
    
    

    for i in MTTR_report:
        
        TOU_MTTRF = i.get('Total Unit')
        Hours_Fail_MTTRF = i.get('Hours Maintenance with Failure is Yes')
        COF_MTTRF = i.get('Count of Failures')
        DBF_MTTRF = i.get('Detail Breakdown with Failure is Yes')
        detail_breakdown_MTTRF = '\n\n'.join(map(str,DBF_MTTRF))
        
        #join([c for c in DBF_MTTRF[0:]])
            
        comment_MTTR = f'Number of Units = {TOU_MTTRF} \nHours maintenance with is Failure Yes ={Hours_Fail_MTTRF} \nCount of failures/Is Failure Yes = {COF_MTTRF} \n\n\n Detail breakdown with failure is yes:\n'
        
        MTTR_value = str(i.get('MTTR'))
        
        row_cells = table_MTTR.add_row().cells
        row_cells[0].text = i.get('Fleet')
        row_cells[1].text = f'MTTRF = {MTTR_value}'
        row_cells[2].text = '6 Hours'
        
        row_cells[3].add_paragraph(comment_MTTR)
        for j in DBF_MTTRF:
            row_cells[3].add_paragraph(j,style='List Bullet')
            

    set_col_widths(table_MTTR)
    Anomali_report_full.add_page_break()

    '''Preventative Maintenance'''
    Anomali_report_full.add_paragraph(text=f'Fleet with PM Below Target {site}').alignment = WD_ALIGN_PARAGRAPH.CENTER


    '''
    formatting PM table
    '''

    table_PM = Anomali_report_full.add_table(rows=1,cols=4)
    table_PM.style = 'Table Grid'



    hdr_cells_PM = table_PM.rows[0].cells
    for cell in hdr_cells_PM:
        cell.height = Mm(30)
        
    hdr_cells_PM[0].text = 'Fleet/Model'
    hdr_cells_PM[1].text = f'PM week {nPeriod}'
    hdr_cells_PM[2].text = 'Target'
    hdr_cells_PM[3].text = 'Comments on Non-Compliance'

    '''cells shading'''
    hdr_cells_PM[0]._tc.get_or_add_tcPr().append(shading_elm13)
    hdr_cells_PM[1]._tc.get_or_add_tcPr().append(shading_elm14)
    hdr_cells_PM[2]._tc.get_or_add_tcPr().append(shading_elm15)
    hdr_cells_PM[3]._tc.get_or_add_tcPr().append(shading_elm16)

    '''repeat header'''
    set_repeat_table_header(table_PM.rows[0])


    for i in PM_report:
        
        TOU_PM = i.get('Total Unit')
        PMD_PM = i.get('PMD total hours')
        UMD_PM = i.get('UMD total hours')
        PDAM_PM = i.get('PDAM total hours')
        
        row_cells_PM = table_PM.add_row().cells
        row_cells_PM[0].text = i.get('Fleet')
        row_cells_PM[1].text = "{:.2%}".format(i.get('PM'))
        row_cells_PM[2].text = "60%"
        
        sch = i.get('Scheduled maintenance')
        Unsch = i.get('Unscheduled maintenance')
        
        Scheduled_maintenance = '\n'.join([c for c in sch[1:]])
        Unscheduled_maintenance = '\n'.join([c for c in Unsch[1:]])
        
        non_compliance = f'Number of Units = {TOU_PM} \nPMD Total = {PMD_PM} \nUMD Total= {UMD_PM} \nPDAM Total hours= {PDAM_PM}\n'
        #MA_anomalies = f'Scheduled maintenance :\n{Scheduled_maintenance}\n\nUnscheduled maintenance :\n{Unscheduled_maintenance}'

        row_cells_PM[3].add_paragraph(text=non_compliance)
        row_cells_PM[3].add_paragraph('Scheduled maintenance :\n')
        for j in sch:
            row_cells_PM[3].add_paragraph(j,style='List Bullet')
        row_cells_PM[3].add_paragraph('\nUnscheduled maintenance :\n')
        for k in Unsch:
            row_cells_PM[3].add_paragraph(k,style='List Bullet')    

    set_col_widths(table_PM)
    Anomali_report_full.add_page_break()


    

    '''Pareto report conditional'''

    if add_pareto_report == True:
        
        '''processing pareto'''

        critical_units_true = fleet_desc_pivot[fleet_desc_pivot['Critical Unit ?'].notna()]
        grouping_for_desc_fleetdesc_pair = All_KPI.groupby(['Description','fleet desc']).size().reset_index()
        grouping_for_desc_fleetdesc_pair.drop(0,axis=1,inplace=True)
        grouping_for_desc_fleetdesc_pair['Critical Unit ?'] = grouping_for_desc_fleetdesc_pair['Description'].apply(lambda x: 'Yes' if x in critical_units_true['Description'].values else np.nan)
        unfiltered_pareto_df = pd.concat([grouping_for_desc_fleetdesc_pair.reset_index(),fleet_desc_pivot.reset_index()])
        pareto_units = unfiltered_pareto_df[unfiltered_pareto_df['Critical Unit ?'].notna()]
        bad_Pareto_this_week = complete_KPI['MA'][complete_KPI['MA']<0.9].dropna()
        Pareto_df = PDTD_complete[PDTD_complete['fleet desc'].isin(pareto_units['fleet desc'])]
        red_Pareto_this_week=bad_Pareto_this_week[bad_Pareto_this_week.index.isin(pareto_units['fleet desc'].tolist())]
        red_Pareto_this_week



        '''Pareto Report Template Column'''
        Pareto_report_fail_df = Pareto_df.groupby(['Unit','Event\nID','Activity','fleet desc','Period\nEvent\nMaint.\nDuration','Is Failure'])['Description of Repair'].apply(lambda x: ','.join(x.astype(str))).reset_index()
        Pareto_report_fail_df['Pareto Anomali'] = Pareto_report_fail_df.loc[:,('Unit')]+' '+Pareto_report_fail_df.loc[:,('Description of Repair')]+' ('+Pareto_report_fail_df.loc[:,('Period\nEvent\nMaint.\nDuration')].values.astype(str)+" Hours)"
        Pareto_report_fail_df.replace({'Pareto Anomali':{'\n':' '}},regex=True)

        Pareto_report = []
 
    
        for i in red_Pareto_this_week.index:
            Downtime_detail_sched = []
            Downtime_detail_unsched = []
            
            Pareto_df_sched = Pareto_report_fail_df[((Pareto_report_fail_df.Activity =='08-Planned Maintenance (PMD)')&(Pareto_report_fail_df['fleet desc'] ==i)&(Pareto_report_fail_df['Period\nEvent\nMaint.\nDuration']>1.0))].sort_values(['Period\nEvent\nMaint.\nDuration'],ascending=False)
            Pareto_df_unsched = Pareto_report_fail_df[((Pareto_report_fail_df.Activity =='09-Unplanned Maintenance (UMD)')&(Pareto_report_fail_df['fleet desc'] ==i)&(Pareto_report_fail_df['Period\nEvent\nMaint.\nDuration']>1.0))].sort_values(['Period\nEvent\nMaint.\nDuration'],ascending=False)
            

            for x in Pareto_df_sched['Pareto Anomali']:
                Downtime_detail_sched.append(x)
            
            for x in Pareto_df_unsched['Pareto Anomali']:
                Downtime_detail_unsched.append(x)

                
            structured_Pareto = {
            'Fleet':i,
            'Pareto':red_Pareto_this_week[i],
            'Total Unit':All_KPI.groupby('fleet desc').count()['Unit'][i],
            'UMD total hours':round(All_KPI.groupby('fleet desc').sum()['UMD\n(09)'][i],2),
            'PMD total hours':All_KPI.groupby('fleet desc').sum()['PMD\n(08)'][i],
            'PDAM total hours':All_KPI.groupby('fleet desc').sum()['PDAM\n08020\n09020'][i],    
                
            'Scheduled maintenance' : Downtime_detail_sched,
            'Unscheduled maintenance' : Downtime_detail_unsched
                
            
        }
            Pareto_report.append(structured_Pareto)

            '''Create word document using word'''
        Pareto_report_full = Document()

        section =Pareto_report_full.sections[0]
        section.page_height = Mm(420)
        section.page_width = Mm(297)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(10.5)
        section.top_margin = Mm(22.2)
        section.bottom_margin = Mm(10.5)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)


        '''Mechanical Availability'''


        '''
        formatting MA table
        '''

        table_Pareto = Pareto_report_full.add_table(rows=1,cols=3)
        table_Pareto.style = 'Table Grid'


        hdr_cells_Pareto = table_Pareto.rows[0].cells
        for cell in hdr_cells_MA:
            cell.height = Mm(30)
            
        hdr_cells_Pareto[0].text = 'Fleet/Model'
        hdr_cells_Pareto[1].text = 'MA this week'
        hdr_cells_Pareto[2].text = 'Comments on Non-Compliance'




        for i in Pareto_report:
            
            TOU_MA = i.get('Total Unit')
            PMD_MA = i.get('PMD total hours')
            UMD_MA = i.get('UMD total hours')
            PDAM_MA = i.get('PDAM total hours')
            
            row_cells_MA = table_Pareto.add_row().cells
            row_cells_MA[0].text = i.get('Fleet')
            row_cells_MA[1].text = "{:.2%}".format(i.get('Pareto'))
            
            sch = i.get('Scheduled maintenance')
            Unsch = i.get('Unscheduled maintenance')
            
            # Scheduled_maintenance = '\n'.join([c for c in sch[1:]])
            # Unscheduled_maintenance = '\n'.join([c for c in Unsch[1:]])
            
            non_compliance = f'Number of Units = {TOU_MA} \nPMD Total = {PMD_MA} \nUMD Total= {UMD_MA} \nPDAM Total hours= {PDAM_MA}\n'
            #MA_anomalies = f'Scheduled maintenance :\n{Scheduled_maintenance}\n\nUnscheduled maintenance :\n{Unscheduled_maintenance}'

            row_cells_MA[2].add_paragraph(text=non_compliance)
            row_cells_MA[2].add_paragraph('Scheduled maintenance :\n')
            for j in sch:
                row_cells_MA[2].add_paragraph(j,style='List Bullet')
            row_cells_MA[2].add_paragraph('\nUnscheduled maintenance :\n')
            for k in Unsch:
                row_cells_MA[2].add_paragraph(k,style='List Bullet')    

        set_col_widths(table_MA)
        Pareto_report_full.add_page_break()

        try:
            save_directory
        except:
            messagebox.showerror('Directory not found', 'Please choose directory to save file')
            return
        else:
            Anomali_report_full.save(os.path.join(save_directory,filename))
            Pareto_report_full.save(os.path.join(save_directory,file_pareto_name))
            messagebox.showinfo(title=None,message=f'report {filename} & {file_pareto_name} done !')
    if add_pareto_report == False:
        try:
            save_directory
        except:
            messagebox.showerror('Directory not found', 'Please choose directory to save file')
            return
        else:    
            Anomali_report_full.save(os.path.join(save_directory,filename))
            messagebox.showinfo(title=None,message=f'report {filename} done !')
            
            
    
    del Pareto_report_full

    del Anomali_report_full





width_button = 18
height_button = 1
padx_button = 0
pady_button = 2


MER_button = Button(window, text ="Upload MER", command=browse_for_file_MER, width=width_button, heigh = height_button, padx = padx_button,pady=pady_button,anchor='e')
EMPS_button = Button(window, text ="Upload EMPS", command=browse_for_file_EMPS, width=width_button, heigh = height_button, padx = padx_button,pady=pady_button,anchor='e')
PDTD_button = Button(window, text ="Upload Downtime Detail", command = browse_for_file_PDTD,width=width_button, heigh = height_button, padx = padx_button,pady=pady_button,anchor='e')
FleetDesc_button = Button(window, text ="Upload Fleet Description", command =browse_for_file_fleetdesc,width=width_button, heigh = height_button, padx = padx_button,pady=pady_button,anchor='e')
where_to_save_file = Button(directory_to_save_frame,text="Save to", command= browse_for_save_location,width=20,height=height_button,pady=pady_button,padx=padx_button,anchor='center')

generate_below_target_report_button =Button(window,text='Generate Report',command=generate_docx, width=14, height=2,font=(5))

developer_label = Label(window,padx = 5, pady=30,text = "Thiess Graduate Development Program, Batch 2021 - Project Improvement\n© Imamf,1108134 \nAny issue regarding this app please contact IFahreza@thiess.com.au")

developer_label.grid(column=1,row=16)



#Welcome.grid(column = 0, row = 0,columnspan=7)
Week_name_label.grid(column=0,row=1)
Week_name_entry.grid(column=1,row=1)
Site_name_Label.grid(column=0,row=2)
Site_name_entry.grid(column=1,row=2)
MER_button.grid(column=0,row=5)
EMPS_button.grid(column=0,row=6)
PDTD_button.grid(column=0,row=7)
FleetDesc_button.grid(column=0,row=8)
where_to_save_file.grid(column=0,row=0)
#MTBF_target_button.grid(column=0,row=9)
#downtime_limit_label.grid(column=0,row=11)

#chk_critical.grid(column=0,row=12)
#chk_pareto.grid(column=0,row=13)

generate_below_target_report_button.grid(column=2,row=15)

window.mainloop()





